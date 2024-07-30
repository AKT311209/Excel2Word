import pandas as pd
import openpyxl as excel
import docx as docu


def readexcel(path, sheetname, range_str):
    """ Reads data from an Excel sheet within a specified range. """
    wb = excel.load_workbook(path, data_only=True,
                             read_only=True, )  # Load the workbook
    sheet = wb[sheetname]  # Access the sheet by name

    # Extract data from the specified range
    data = []
    for row in sheet.iter_rows(min_row=int(range_str.split(':')[0][1:]),
                               max_row=int(range_str.split(':')[1][1:]),
                               min_col=ord(range_str.split(':')[0][0]) - 64,
                               max_col=ord(range_str.split(':')[1][0]) - 64):
        row_data = []
        for cell in row:
            if cell.value is None:
                row_data.append(None)
            else:
                row_data.append(cell.value)
        data.append(row_data)

    df = pd.DataFrame(data)
    return df


def create_table_in_docx(data, filename):
    doc = docu.Document(filename)

    # Check if a table exists
    if doc.tables:
        table = doc.tables[0]
        # Modify the table data while preserving formatting
        for i, row in enumerate(data.itertuples(index=False)):
            for j, cell_value in enumerate(row):
                # Get the existing paragraph in the cell
                paragraph = table.cell(i, j).paragraphs[0]
                # Check if the cell is empty
                if not paragraph.runs:
                    # Add a new run with the cell value and copy formatting
                    new_run = paragraph.add_run(str(cell_value))
                    # Copy formatting from surrounding cells
                    if i > 0 and j > 0:
                        # Top-left cell
                        previous_cell_paragraph = table.cell(
                            i - 1, j - 1).paragraphs[0]
                        if previous_cell_paragraph.runs:
                            copy_formatting(
                                new_run, previous_cell_paragraph.runs[0])
                    elif i > 0:
                        # Above cell
                        previous_cell_paragraph = table.cell(
                            i - 1, j).paragraphs[0]
                        if previous_cell_paragraph.runs:
                            copy_formatting(
                                new_run, previous_cell_paragraph.runs[0])
                    elif j > 0:
                        # Left cell
                        previous_cell_paragraph = table.cell(
                            i, j - 1).paragraphs[0]
                        if previous_cell_paragraph.runs:
                            copy_formatting(
                                new_run, previous_cell_paragraph.runs[0])
                else:
                    # Replace the existing text with the new text
                    paragraph.runs[0].text = str(
                        cell_value)  # Update the first run
    else:
        # Create a new table
        table = doc.add_table(rows=len(data) + 1, cols=len(data.columns))
        # Populate the new table
        for i, row in enumerate(data.itertuples(index=False)):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)
    doc.save(filename)


def copy_formatting(new_run, source_run):
    """Copies formatting from a source run to a new run."""
    new_run.font.name = source_run.font.name
    new_run.font.size = source_run.font.size
    new_run.font.bold = source_run.font.bold
    new_run.font.italic = source_run.font.italic
    new_run.font.underline = source_run.font.underline
    new_run.font.strike = source_run.font.strike
    new_run.font.color.rgb = source_run.font.color.rgb


if __name__ == "__main__":
    sheet_name = input("Enter the sheet name: ")
    range_sheet = input("Enter the range (e.g., A1:C5): ")
    doc_name = input("Enter the output file name: ")
    excel_input = input("Enter the name of the Excel file: ")
    if doc_name:
        doc_name = doc_name + ".docx"
    if excel_input:
        excel_input = excel_input + ".xlsx"

    datafromex = readexcel(excel_input, sheet_name, range_sheet)
    create_table_in_docx(datafromex, doc_name)

    print("Table created successfully!")
